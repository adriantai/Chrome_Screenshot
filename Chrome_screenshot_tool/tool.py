from docx import Document
from selenium import webdriver
from selenium.webdriver.common.keys import Keys
from selenium.webdriver.common.by import By
import time
import pandas as pd
from docx.shared import Inches
from PIL import Image
from typing import List, Dict, Any, Optional
import PySimpleGUI as sg
import os

Chinese_MAX = 4
English_MAX = 3


class No_page_error(Exception):
    """
    Raised when there is no "next" button on a google search page
    """
    print("Next page not found")


def word_counter(word: str, language: str) -> int:
    """
    Only use this for counting names.
    Counts how many words are in the string word.
    Works for both Chinese and English
    DOES NOT work when both chinese and english are mixed and not separated
    by an empty space
    Returns -1 on error
    >>> word_counter("hello darkness", 'E')
    2
    >>> word_counter("Apple", 'E')
    1
    >>> word_counter("hello darkness my good friend", 'E')
    5
    >>> word_counter('德张民智开明范 学领女权炳耀风', 'C')
    14
    >>> word_counter("实验室实验室", 'C')
    6
    """
    if language == 'E':
        return len(word.split(" "))
    if language == 'C':
        t = word.split(" ")
        c = 0
        for x in range(len(t)):
            c += len(t[x])
        return c
    return -1


def import_pic(image: str, doc_path: str, one_page: bool) -> bool:
    """
    Imports image into a word document directed by doc_path.
    preconditions:
    - doc_path must be a string ending in .docx
    - image must be a string ending in .png
    """
    doc = Document(doc_path)
    section = doc.sections[0]
    h = section.page_height.inches - section.bottom_margin.inches
    # see https://python-docx.readthedocs.io/en/latest/user/sections.html

    if one_page is False:
        doc.add_picture(image, width=Inches(6.5))
    else:
        open_image = Image.open(image)
        open_image.thumbnail((1920, 1080))
        open_image.save('img.png')
        doc.add_picture(open_image)
    doc.save(doc_path)
    return True


def screenshot_one(search: str, pages: int, doc: str) -> bool:
    """
    Takes one full page of google search query given search.
    If possible, it redirects to page 2 (and beyond) of the google search,
    and takes another full page screenshot.
    Precondition:
        pages > 0
        doc must be a valid word document name (without .docx)
    Zoom level is assumed to be 70% by default
    Returns true when all pages are screenshotted, returns false when not all
    google search pages specified by pages are not screenshotted
    i.e. pages > google serach pages

    >>> screenshot_one("dog", 1, "template")
    True
    """
    options = webdriver.ChromeOptions()
    options.add_argument('--headless=new')  # headless hides chrome UI
    options.add_argument("--lang=en-GB")
    driver = webdriver.Chrome(options=options)
    driver.set_window_size(1920, 4320)

    driver.implicitly_wait(3)
    driver.get('https://www.google.com')  # navigating to google
    find = driver.find_element(By.NAME, 'q')

    checked_search = ""
    alphabet = "qwertyuiopasdfghjklzxcvbnmQWERTYUIOPASDFGHJKLZXCVBNM"
    for i in alphabet:
        if search[0] not in i:  # e.g chinese name
            checked_search += "'" + search + "'"
            break
    if search[0] in alphabet:
        checked_search = "(" + search + ")"

    find.send_keys(checked_search)
    find.send_keys(Keys.RETURN)

    links_count = pages
    curr_ss = 0
    normal_exit = True

    while links_count != 0:
        # driver.execute_script("document.body.style.zoom='50%'")
        time.sleep(0.4)
        links_count -= 1
        # print("current page at " + str((links_count - pages) + 1))

        # Cap screenshot
        ss_name_top = "ss" + str(curr_ss) + ".png"
        driver.save_screenshot(ss_name_top)
        height = driver.execute_script("return document.")
        import_pic(ss_name_top, doc + ".docx", True)
        curr_ss += 1

        driver.execute_script("document.body.style.zoom='100%'")
        time.sleep(0.4)
        next_links = driver.find_elements(By.LINK_TEXT, "Next")
        if len(next_links):
            # print('Found "Next" link')
            next_links[0].click()
        else:
            print('There is no "Next" link')
            normal_exit = False
            driver.quit()
    # Delete screenshots
    for i in range(curr_ss):
        to_remove = os.getcwd() + '\\' + 'ss' + str(i) + '.png'
        os.remove(to_remove)
    return normal_exit


def screenshot_two(search: str, pages: int, doc: str, language: str) -> bool:
    """
    Takes two screenshots per google search query given search. The
    first screenshot captures the first half of the google search, and the
    second screenshot captures the second half.
    If possible, it redirects to page 2 (and beyond) of the google search,
    and takes two more (top & bottom half) screenshots.
    Precondition:
        pages > 0
        doc must be a valid word document name (without .docx)
    Zoom level is assumed to be 57% by default
    Returns true when all pages are screenshotted, returns false when not all
    google search pages specified by pages are not screenshotted
    i.e. pages > google serach pages
    >>> screenshot_two("cat", 2, "template", 'E')
    True
    """
    options = webdriver.ChromeOptions()
    options.add_argument('--headless=new')  # headless hides chrome UI
    # options.add_argument("--lang=en-GB")
    driver = webdriver.Chrome(options=options)
    driver.set_window_size(1920, 1200)

    driver.implicitly_wait(3)
    driver.get('https://www.google.com.hk')  # navigating to google
    find = driver.find_element(By.NAME, 'q')
    checked_search = ""
    if language == 'E':
        checked_search = search
    if language == 'C':
        checked_search = "'" + search + "'"

    find.send_keys(checked_search)
    find.send_keys(Keys.RETURN)

    links_count = pages
    curr_ss = 0
    normal_exit = True
    while links_count != 0:
        driver.execute_script("document.body.style.zoom='70%'")
        time.sleep(0.4)
        links_count -= 1
        # print("current page at " + str((links_count - pages) + 1))

        # Cap top screenshot
        ss_name_top = "ss" + str(curr_ss) + ".png"
        driver.save_screenshot(ss_name_top)
        import_pic(ss_name_top, doc + ".docx", False)
        curr_ss += 1

        # Scroll and cap bottom screenshot
        driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
        time.sleep(0.3)
        ss_name_bottom = "ss" + str(curr_ss) + ".png"
        driver.save_screenshot(ss_name_bottom)

        import_pic(ss_name_bottom, doc + ".docx", False)
        curr_ss += 1
        driver.execute_script("document.body.style.zoom='100%'")
        time.sleep(0.4)
        next_links = driver.find_elements(By.LINK_TEXT, "下一頁")
        # next_links = driver.find_elements(By.LINK_TEXT, "Next")
        if len(next_links):
            # print('Found "Next" link')
            next_links[0].click()
        else:
            print('There is no "Next" link')
            normal_exit = False
            driver.quit()
    # Delete screenshots
    for i in range(curr_ss):
        to_remove = os.getcwd() + '\\' + 'ss' + str(i) + '.png'
        os.remove(to_remove)
    return normal_exit


def split_prompt(prompt: str, language: str) -> List[str]:
    """
    splits a prompt into roughly two halves
    >>> split_prompt("AND (bunny OR dog OR turtle OR cat OR mouse)", 'E')
    ['AND (bunny OR dog)', 'AND (turtle OR cat OR mouse)']
    >>> split_prompt("AND (‘兔’OR‘狗’OR‘猫’OR‘马’)", 'C')
    ['AND ('兔’OR'狗’)', 'AND (‘猫’OR‘马’)']
    >>> p = 'AND (“攀冰”OR“攀岩”OR“跳傘”OR“釣魚”OR“鬥雞”OR“賽馬”OR“跳遠”OR“跳高”OR“競走”OR“走路”OR“滑板”OR“單板滑雪”OR“衝浪”OR“棒術”OR“忍術”OR“柔道”OR“拳擊”OR“泰拳”OR“排球”OR“足球”OR“滑水”OR“彈床”OR“賽車”OR“跑酷”OR“乒乓球”OR“拔河”OR“象棋”)'
    >>> split_prompt(p, 'C')
    ['AND (“攀冰”OR“攀岩”OR“跳傘”OR“釣魚”OR“鬥雞”OR“賽馬”OR“跳遠”OR“跳高”OR“競走”OR“走路”OR“滑板”OR“單板滑雪”OR“衝浪”)', 'AND (“棒術”OR“忍術”OR“柔道”OR“拳擊”OR“泰拳”OR“排球”OR“足球”OR“滑水”OR“彈床”OR“賽車”OR“跑酷”OR“乒乓球”OR“拔河”OR“象棋”)']
    """
    #  Unsimplified code
    close_bracket = len(prompt)
    opening_bracket = 0
    for i in range(len(prompt)):
        if prompt[i] == '(':
            break
        opening_bracket += 1
    prompt_list = prompt[opening_bracket + 1:close_bracket - 1]
    keywords = prompt_list.split(" OR ")
    mid_index = len(keywords) // 2
    first_list = 'AND ('
    second_list = 'AND ('
    if language == 'E':
        for i in range(mid_index):
            if i+1 == mid_index:
                first_list = first_list + keywords[i] + ')'
            else:
                first_list = first_list + keywords[i] + ' OR '
        for i in range(len(keywords) - mid_index):
            x = i + mid_index
            if x + 1 == len(keywords):
                second_list = second_list + keywords[x] + ')'
            else:
                second_list = second_list + keywords[x] + ' OR '
        return [first_list, second_list]
    if language == 'C':
        Chinese_words = prompt[opening_bracket + 1:len(prompt) - 1]
        Chinese_words = Chinese_words.split('OR')

        mid_index = len(Chinese_words) // 2
        for i in range(mid_index):

            if i+1 == mid_index:
                first_list += Chinese_words[i] + ')'

            else:
                first_list = first_list + Chinese_words[i] + 'OR'
        for i in range(len(Chinese_words) - mid_index):
            x = i + mid_index
            if x + 1 == len(Chinese_words):
                second_list = second_list + Chinese_words[x] + ')'
            else:
                second_list = second_list + Chinese_words[x] + 'OR'
        return [first_list, second_list]


def execute_search(keyword: str, num_ss: int, language: str, doc: str,
                   prompt: str):
    """
    doc must be a valid word document name or path without ".docx"
    num_ss is the number of screenshots desired to take
    language can either be 'E' for English or 'C' for Chinese
    prompt is the prompt used after the keyword. prompt must be in the format of
    "AND (A OR B OR C)".
    >>> e = 'AND ("Apple" OR "Avocado" OR "Blueberry" OR "Cacao" OR "Date" OR "Cherry" OR "Coconut" OR "Fig" OR "Dragonfruit" OR "Durian" OR "Grape" OR "Raisin" OR "Plum" OR "Mango" OR "Melon" OR "Orange" OR "Strawberry" OR "Jackfruit" OR "Guava" OR "Yuzu" OR "Kiwano" OR "Lime" OR "Lemon" OR "Honeyberry" OR "Peach" OR "Pear" OR "Cranberry")'
    >>> m = 'AND (“攀冰”OR“攀岩”OR“跳傘”OR“釣魚”OR“鬥雞”OR“賽馬”OR“跳遠”OR“跳高”OR“競走”OR“走路”OR“滑板”OR“單板滑雪”OR“衝浪”OR“棒術”OR“忍術”OR“柔道”OR“拳擊”OR“泰拳”OR“排球”OR“足球”OR“滑水”OR“彈床”OR“賽車”OR“跑酷”OR“乒乓球”OR“拔河”OR“象棋”)'
    >>> execute_search("中国建設銀行", 2, 'C', "template", m)
    """
    # code is not simplified to make it easier to read
    if language == 'E':
        keyword = "(" + keyword + ")"
        if len(prompt) == 0:
            screenshot_two(keyword, num_ss, doc, 'E')
        elif word_counter(keyword, 'E') <= English_MAX:
            screenshot_two((keyword + " " + prompt), num_ss, doc, 'E')
        else:
            prompt_split = split_prompt(prompt, 'E')
            screenshot_two((keyword + " " + prompt_split[0]), num_ss, doc, 'E')
            screenshot_two((keyword + " " + prompt_split[1]), num_ss, doc, 'E')

    if language == 'C':
        keyword = keyword + "'"
        if len(prompt) == 0:
            screenshot_two(keyword, num_ss, doc, 'C')
        elif word_counter(keyword, 'C') <= Chinese_MAX:
            screenshot_two((keyword + " " + prompt), num_ss, doc, 'C')
        else:
            prompt_split = split_prompt(prompt, 'C')
            screenshot_two((keyword + " " + prompt_split[0]), num_ss, doc, 'C')
            screenshot_two((keyword + " " + prompt_split[1]), num_ss, doc, 'C')


class Handler:
    Excel_path: str
    Excel_sheet: str
    Folder_path: str
    Chinese_prompt: str
    English_prompt: str
    pages: int

    def __init__(self):
        self.Excel_path = ""
        self.Folder_path = ""
        self.Excel_sheet = ""
        self.Chinese_prompt = ""
        self.English_prompt = ""
        self.pages = 1

    def get_excel_path(self):
        print(self.Excel_path)

    def get_excel_sheet(self):
        print(self.Excel_sheet)

    def get_folder_path(self):
        print(self.Folder_path)

    def get_english_prompt(self):
        print(self.English_prompt)

    def get_chinese_prompt(self):
        print(self.Chinese_prompt)

    def handle_excel(self):
        excel = pd.read_excel(self.Excel_path, sheet_name=self.Excel_sheet)
        columns = excel.columns.ravel()

        Col0_values = excel[columns[0]].tolist()
        Col1_values = excel[columns[1]].tolist()
        Col2_values = excel[columns[2]].tolist()

        row = len(Col0_values)
        for i in range(row):
            folder_name = str(Col2_values[i]) + "_" + Col0_values[i]
            # folder_directory = os.getcwd() + '\\' + folder_name
            template_dir = self.Folder_path + '\\' + folder_name + '\\' + \
                           "template"
            if str(Col0_values[i]) != "nan":
                execute_search(Col0_values[i], self.pages, 'E', template_dir,
                               self.English_prompt)  # English
            if str(Col1_values[i]) != "nan":
                execute_search(Col1_values[i], self.pages, 'C', template_dir,
                               self.Chinese_prompt)  # Chinese

    def handle_gui(self):
        """
        GUI for this class
        """
        developer = [sg.Checkbox("Developer mode", key="-dev-", default=False)]
        excel_path = [sg.Text("Select Excel file"),
                      sg.Input(key="-excel_path-"), sg.FileBrowse()]
        excel_sheet = [sg.Text("Enter Excel sheet name"),
                       sg.InputText(key="-excel_sheet-")]

        Eng_prompt = [sg.Text("Enter the English prompt"),
                      sg.Multiline(size=(60, 4), key='-Eng_prompt-')]
        Chi_prompt = [sg.Text("Enter the Chinese prompt"),
                      sg.Multiline(size=(60, 4), key='-Chi_prompt-')]

        page_num = [sg.Text("How many pages of Google search? (Number only)"),
                    sg.InputText(key="-page_num-")]

        b1 = sg.Button("Submit", font=('Arial Bold', 10))

        layout = [[sg.Text("Google search demo")], excel_path,
                  excel_sheet, Eng_prompt, Chi_prompt, page_num, [b1, developer]]

        # Create the window
        window = sg.Window(title="Google search Demo", layout=layout,
                           margins=(100, 100))

        # Create an event loop
        while True:
            event, values = window.read()
            self.Excel_path = values["-excel_path-"]
            self.Excel_sheet = values["-excel_sheet-"]
            excel_name = self.Excel_path.split("/")[-1]
            truncate = len(self.Excel_path) - len(excel_name)
            self.Folder_path = self.Excel_path[:truncate]
            self.English_prompt = values["-Eng_prompt-"]
            self.Chinese_prompt = values["-Chi_prompt-"]
            self.pages = int(values["-page_num-"])
            # End program if user closes window or
            # presses the OK button
            if event == sg.WIN_CLOSED:
                break
            if event == "Submit":
                # sg.popup("Click OK to confirm.")
                if values["-dev-"] is False:
                    self.handle_excel()
                else:
                    pass
                break
        window.close()


if __name__ == '__main__':
    Handler().handle_gui()

"""
    try:
        driver.implicitly_wait(3)
        driver.get('https://www.google.com')  # navigating to google
        search_gog = driver.find_element(By.NAME, 'q')
        search_gog.send_keys(search)
        search_gog.send_keys(Keys.RETURN)
        time.sleep(1)
        next_links = driver.find_elements(By.LINK_TEXT, "Next")
        if len(next_links):
            print('Found "Next" link')
            next_links[0].click()
        else:
            print('There is no "Next" link')
    finally:
        input('pausing (hit enter to terminate) ...')
        driver.quit()
"""

"""
def screenshot_test(search: str, pages: int) -> None:
    
options = webdriver.ChromeOptions()
options.add_argument('--headless=new')  # headless hides chrome UI
options.add_argument("--lang=en-GB")
driver = webdriver.Chrome(options=options)
driver.set_window_size(1920, 1200)

driver.implicitly_wait(3)
driver.get('https://www.google.com')  # navigating to google
find = driver.find_element(By.NAME, 'q')

find.send_keys(search)
find.send_keys(Keys.RETURN)
# driver.find_elements(By.LINK_TEXT, "Next")
links_count = pages
ss_count = links_count * 2  # number of screenshots to be taken at most
curr_ss = 1
while links_count != 0:
    driver.execute_script("document.body.style.zoom='70%'")
    time.sleep(0.3)
    links_count -= 1
    print("current page at " + str((links_count - pages) + 1))

    # Cap top screenshot
    ss_name_top = "ss" + str(curr_ss) + ".png"
    driver.save_screenshot(ss_name_top)
    # import_pic(ss_name_top, "template.docx")
    curr_ss += 1

    # Scroll and cap bottom screenshot
    # driver.execute_script("window.scrollTo(0, window.scrollY + 800)")
    driver.execute_script("window.scrollTo(0, document.body.scrollHeight);")
    time.sleep(0.2)
    ss_name_bottom = "ss" + str(curr_ss) + ".png"
    driver.save_screenshot(ss_name_bottom)

    # import_pic(ss_name_bottom, "template.docx")
    curr_ss += 1
    driver.execute_script("document.body.style.zoom='100%'")
    time.sleep(0.2)
    next_links = driver.find_elements(By.LINK_TEXT, "Next")
    if len(next_links):
        print('Found "Next" link')
        next_links[0].click()
    else:
        print('There is no "Next" link')
        driver.quit()

"""
# def handle_excel():
#     """
#     >>> handle_excel()
#
#     """
#     sheet = "target"
#     excel = pd.read_excel("testBook.xlsx", sheet_name=sheet)
#     columns = excel.columns.ravel()
#
#     Col0_values = excel[columns[0]].tolist()
#     Col1_values = excel[columns[1]].tolist()
#     Col2_values = excel[columns[2]].tolist()
#
#     row = len(Col0_values)
#     for i in range(row):
#         print(Col1_values[i])
#         if str(Col1_values[i]) == "nan":
#             print("True")
#     for i in range(row):
#         folder_name = str(Col2_values[i]) + "_" + Col0_values[i]
#         folder_directory = os.getcwd() + '\\' + folder_name
#         template_dir = os.getcwd() + '\\' + folder_name + '\\' + "template"
#         if str(Col0_values[i]) != "nan":
#             execute_search(Col0_values[i], 1, 'E', template_dir, "")  # English
#         if str(Col1_values[i]) != "nan":
#             execute_search(Col1_values[i], 1, 'C', template_dir, "")  # Chinese


# Prompt

"""
AND ("Apple" OR "Avocado" OR "Blueberry" OR "Cacao" OR "Date" OR "Cherry" OR "Coconut" OR "Fig" OR "Dragonfruit" OR "Durian" OR "Grape" OR "Raisin" OR "Plum" OR "Mango" OR "Melon" OR "Orange" OR "Strawberry" OR "Jackfruit" OR "Guava" OR "Yuzu" OR "Kiwano" OR "Lime" OR "Lemon" OR "Honeyberry" OR "Peach" OR "Pear" OR "Cranberry")
AND (“攀冰”OR“攀岩”OR“跳傘”OR“釣魚”OR“鬥雞”OR“賽馬”OR“跳遠”OR“跳高”OR“競走”OR“走路”OR“滑板”OR“單板滑雪”OR“衝浪”OR“棒術”OR“忍術”OR“柔道”OR“拳擊”OR“泰拳”OR“排球”OR“足球”OR“滑水”OR“彈床”OR“賽車”OR“跑酷”OR“乒乓球”OR“拔河”OR“象棋”)
"""
